from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "public/brand/cnvrt-logo.png"
TEMPLATE = ROOT / "templates/documents/cnvrt-website-proposal-template-v1.docx"
CLIENT = Path("/Users/danlyons/Documents/RejuvinationDoctors/The-Rejuvenation-Doctors-Website-Proposal.docx")

INK = "171719"
MUTED = "66636B"
LAVENDER = "C7B6DA"
LAVENDER_LIGHT = "F4F0F8"
WHITE = "FFFFFF"
LINE = "DDD9E1"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=140, start=180, bottom=140, end=180):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")

def set_run(run, size=10.5, color=INK, bold=False, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

def add_text(doc, text, size=10.5, color=INK, bold=False, italic=False, after=8, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.24
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 18, 2: 13.5, 3: 11.5}[level], color=INK, bold=True)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        set_run(p.add_run(f"  {item}"), size=10.25)

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), LAVENDER)
    borders.append(bottom)
    p_pr.append(borders)

def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.24
    for level, size, before, after in ((1, 18, 16, 8), (2, 13.5, 12, 6), (3, 11.5, 8, 4)):
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(10.25)
    bullet.paragraph_format.left_indent = Inches(0.38)
    bullet.paragraph_format.first_line_indent = Inches(-0.19)
    bullet.paragraph_format.space_after = Pt(4)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("CNVRT  /  PRIVATE CLIENT PROPOSAL"), size=8, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    set_run(footer.add_run("CNVRT · Manchester, UK · hello@cnvrtdigital.co.uk"), size=8, color=MUTED)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_cover(doc, client, contact, date, reference):
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Inches(6.9)
    cell = banner.cell(0, 0)
    cell.width = Inches(6.9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, INK)
    set_cell_margins(cell, top=260, start=260, bottom=260, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(LOGO), width=Inches(1.55))
    add_text(doc, "WEBSITE PROJECT PROPOSAL", size=9, color="8063A4", bold=True, after=10)
    add_text(doc, "A clearer digital patient journey\nfor hair restoration and medical aesthetics.", size=25, bold=True, after=12)
    add_text(doc, f"Prepared for {client}", size=13, color=MUTED, after=26)
    add_rule(doc)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.55), Inches(5.25))
    rows = (("Client", client), ("Contact", contact), ("Prepared", date), ("Reference", reference))
    for row, values in zip(table.rows, rows):
        for i, (value, width) in enumerate(zip(values, widths)):
            row.cells[i].width = width
            set_cell_margins(row.cells[i], top=80, start=100, bottom=80, end=100)
            p = row.cells[i].paragraphs[0]
            set_run(p.add_run(value), size=9.5, color=MUTED if i == 0 else INK, bold=i == 0)
    add_text(doc, "Prepared by Daniel Lyons, Founder & Growth Specialist, CNVRT", size=9.5, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

def add_investment_table(doc, total, deposit, balance):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.15)
    table.columns[1].width = Inches(1.65)
    for i, text in enumerate(("Investment", "Amount")):
        cell = table.rows[0].cells[i]
        shade(cell, INK)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i else WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(text), size=9.5, color=WHITE, bold=True)
    rows = (("Website strategy, design, development and launch", total), ("Deposit to schedule the project", deposit), ("Balance due before launch", balance))
    for label, amount in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell)
        shade(cells[0], LAVENDER_LIGHT if len(table.rows) % 2 == 0 else WHITE)
        shade(cells[1], LAVENDER_LIGHT if len(table.rows) % 2 == 0 else WHITE)
        set_run(cells[0].paragraphs[0].add_run(label), size=9.5, bold=label.startswith("Website"))
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run(cells[1].paragraphs[0].add_run(amount), size=9.5, bold=True)
    return table

def build(path, values):
    doc = Document()
    configure(doc)
    add_cover(doc, values["client"], values["contact"], values["date"], values["reference"])
    doc.add_page_break()
    add_heading(doc, "The opportunity", 1)
    add_text(doc, values["opportunity"])
    add_text(doc, "Recommended position", size=9, color="8063A4", bold=True, after=4)
    add_text(doc, values["position"], size=12, bold=True, after=12)
    add_heading(doc, "Project objectives", 2)
    add_bullets(doc, values["objectives"])
    add_heading(doc, "What the project will deliver", 1)
    for heading, paragraph, bullets in values["scope"]:
        add_heading(doc, heading, 2)
        add_text(doc, paragraph)
        if bullets:
            add_bullets(doc, bullets)
    add_heading(doc, "Patient journeys and management", 1)
    add_text(doc, values["journey_intro"])
    for heading, bullets in values["journeys"]:
        add_heading(doc, heading, 2)
        add_bullets(doc, bullets)
    add_heading(doc, "Delivery approach", 1)
    for heading, text in values["process"]:
        add_heading(doc, heading, 3)
        add_text(doc, text)
    add_text(doc, values["timeline"], color=MUTED, italic=True)
    doc.add_page_break()
    add_heading(doc, "Investment", 1)
    add_text(doc, "A fixed project investment covering the agreed phase-one website scope.")
    add_investment_table(doc, values["total"], values["deposit"], values["balance"])
    add_heading(doc, "Included", 2)
    add_bullets(doc, values["included"])
    add_heading(doc, "Not included unless agreed", 2)
    add_bullets(doc, values["excluded"])
    add_heading(doc, "Client responsibilities", 2)
    add_bullets(doc, values["responsibilities"])
    doc.add_page_break()
    add_heading(doc, "Agreement", 1)
    add_text(doc, values["agreement"])
    add_heading(doc, "Key terms", 2)
    add_bullets(doc, values["terms"])
    add_text(doc, "Accepted for and on behalf of the client", size=9, color=MUTED, bold=True, after=14)
    for label in ("Name", "Signature", "Date"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
        set_run(p.add_run(f"{label}: "), size=10, bold=True)
        set_run(p.add_run("____________________________________________________________"), size=10, color=MUTED)
    add_rule(doc)
    add_text(doc, "CNVRT", size=14, bold=True, after=3)
    add_text(doc, "Daniel Lyons · Founder & Growth Specialist\nhello@cnvrtdigital.co.uk · 07742 119 335 · cnvrtdigital.co.uk", size=9.5, color=MUTED, after=0)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)

base = {
    "date": "20 July 2026", "reference": "CNVRT-PROP-TRD-01",
    "opportunity": "The current website contains substantial treatment knowledge and proof, but the patient journey is fragmented across older treatment pages, Wix booking services, inconsistent location signals and multiple contact details. The new website should make the clinic easier to trust, understand and book, while giving the business a maintainable foundation for search visibility and follow-up.",
    "position": "Lead with doctor-led hair restoration. Present medical aesthetics as a strong, clearly organised secondary service line.",
    "objectives": ["Increase qualified consultation bookings, particularly for hair restoration.", "Establish Dr Jamie’s medical credibility, personal experience and natural-results philosophy.", "Give each priority treatment a clear, accurate and commercially useful patient journey.", "Replace inconsistent location, pricing and contact information with one approved source of truth.", "Create an SEO-ready structure for Manchester, Bury and any other confirmed clinic locations.", "Provide a practical enquiry and consultation-management view without storing confidential clinical records."],
    "scope": [
        ("Strategy and structure", "CNVRT will define the website architecture around the questions patients ask before choosing a clinic.", ["Priority treatment and location architecture", "Hair restoration and medical aesthetics navigation", "Conversion paths for consultations, calls, WhatsApp and forms"]),
        ("Design and responsive build", "A bespoke, mobile-first website will be designed and developed in a premium clinical style that feels credible, human and distinctive.", ["Core page design system", "Responsive templates for treatment, practitioner, result and article content", "Accessible interaction and clear calls to action"]),
        ("Content restructuring", "Existing website content will be audited, consolidated and rewritten for clarity, subject to Dr Jamie’s medical approval.", ["Homepage and Dr Jamie narrative", "Priority service and treatment pages", "Trust, safety, consultation and FAQ content", "Initial four SEO-led articles"]),
        ("Migration and launch", "CNVRT will prepare the agreed content, redirects and measurement foundation needed to move away from the current Wix structure.", ["Content and media migration within the agreed scope", "Core on-page SEO and redirect plan", "Analytics and Search Console connection", "Launch checks and handover"]),
    ],
    "journey_intro": "The new system will separate marketing operations from clinical records. It can organise enquiries, consultations and follow-up activity, but confidential medical information must remain in the clinic’s dedicated clinical system.",
    "journeys": [("Patient-facing journey", ["Discover the right treatment", "Understand suitability, process, price guidance and risks", "Choose an online or in-clinic consultation", "Submit an enquiry or booking request", "Receive clear confirmation and next steps"]), ("Admin-facing journey", ["View website enquiries and booking sources", "Move patients through an agreed enquiry pipeline", "Record notes and follow-up actions", "Review booking and conversion activity", "Manage approved website content and patient results within the agreed system scope"])],
    "process": [("01 · Discovery and approval", "Confirm treatments, locations, pricing, booking rules, credentials, access and regulatory information."), ("02 · Structure and content", "Agree the sitemap and patient journeys, then prepare the priority page content for medical review."), ("03 · Design and build", "Design the responsive interface and build the approved page system, booking routes and admin requirements."), ("04 · Review and launch", "Complete content approval, responsive QA, analytics, redirects, access checks and launch handover.")],
    "timeline": "Target delivery sequence: approximately 6–8 weeks from receipt of the deposit, completed questionnaire, required access and timely content approvals. A dated schedule will be confirmed after discovery.",
    "total": "£2,000.00", "deposit": "£1,000.00", "balance": "£1,000.00",
    "included": ["The phase-one strategy, design, development and content scope described in this proposal.", "Two structured review rounds at the design/content stage.", "Responsive and launch-quality checks for current desktop and mobile browsers.", "Thirty days of reasonable post-launch defect support."],
    "excluded": ["Third-party software, hosting, domain, SMS, WhatsApp, email or booking-platform charges.", "Paid advertising spend, ongoing SEO retainers or ongoing content production.", "Clinical record storage or migration into the website CRM.", "New professional photography, video production or legal/regulatory advice.", "Material scope additions requested after approval."],
    "responsibilities": ["Provide complete and accurate clinic, pricing, policy, qualification, GMC and CQC information.", "Supply original media and confirm written patient permission for every result or testimonial used.", "Review medical content and provide consolidated feedback within agreed review windows.", "Provide access to Wix, domain, booking tools, Analytics, Search Console and Google Business Profile.", "Nominate one person with final authority to approve design, content and launch."],
    "agreement": "Approval confirms the project direction, investment and responsibilities described in this proposal. The detailed discovery response and approved sitemap will form the working delivery scope. Any material additions will be estimated and agreed before work proceeds.",
    "terms": ["The £1,000 deposit is required to reserve the project and begin delivery.", "The remaining £1,000 is due before the website is launched or transferred.", "Invoices are payable within seven days unless otherwise agreed in writing.", "CNVRT retains ownership of working files and unpublished work until invoices are paid in full.", "The client is responsible for final medical, legal, regulatory and factual approval.", "Either party may pause the project where required information, access, approval or payment is outstanding."],
}

client_values = {**base, "client": "The Rejuvenation Doctors", "contact": "Dr Jamie"}
template_values = {**base, "client": "{{CLIENT_NAME}}", "contact": "{{CLIENT_CONTACT}}", "date": "{{PROPOSAL_DATE}}", "reference": "{{PROPOSAL_REFERENCE}}", "total": "{{TOTAL_INVESTMENT}}", "deposit": "{{DEPOSIT}}", "balance": "{{BALANCE}}"}
build(CLIENT, client_values)
build(TEMPLATE, template_values)
print(CLIENT)
print(TEMPLATE)
